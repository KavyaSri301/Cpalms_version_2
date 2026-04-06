import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import base64

def generate_docx_file(content: str, title: str = "CPALMS AI Customization"):
    """Generates a DOCX file with the given content"""
    content = re.sub(r'\n\s*\n+', '\n', content.strip())
    
    doc = Document()
    doc.add_heading(title, level=0)
    
    for para in content.split("\n"):
        if not para.strip():
            doc.add_paragraph("")
            continue
        
        paragraph = doc.add_paragraph()
        while "**" in para:
            before, rest = para.split("**", 1)
            if "**" in rest:
                bold_text, after = rest.split("**", 1)
                paragraph.add_run(before)
                run = paragraph.add_run(bold_text)
                run.bold = True
                para = after
            else:
                paragraph.add_run(before + "**" + rest)
                para = ""
        paragraph.add_run(para)
    
    return doc


def make_docx_link(doc_buffer):
    """Convert a DOCX BytesIO buffer to a base64 data URI"""
    doc_buffer.seek(0)
    b64 = base64.b64encode(doc_buffer.read()).decode()
    return f'data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}'


def generate_docx_file_for_download(content: str, title: str = "CPALMS AI Customization"):
    """Generates a DOCX file with a centered blue heading"""
    doc = Document()
    heading = doc.add_paragraph()
    run = heading.add_run("AI Customization")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    content = re.sub(r'<a\s+href="[^"]+"[^>]*>.*?</a>', r'www.cpalms.org', content)
    
    def format_line(line):
        segments = []
        while '**' in line:
            before, rest = line.split('**', 1)
            if '**' not in rest:
                segments.append(('text', before + '**' + rest))
                return segments

            bold_text, after = rest.split('**', 1)
            segments.append(('text', before))
            segments.append(('bold', bold_text))
            line = after

        segments.append(('text', line))
        return segments

    
    blocks = re.split(r'\n\s*\n', content.strip())
    for block in blocks:
        lines = block.strip().splitlines()
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            para = doc.add_paragraph()
            for typ, txt in format_line(line):
                run = para.add_run(txt + ' ')
                if typ == 'bold':
                    run.bold = True
    
    return doc
